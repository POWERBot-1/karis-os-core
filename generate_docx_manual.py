#!/usr/bin/env python3
"""
Generates the professional Word Document (.docx) for KARIS OS™ Architecture & Build Manual.
Run: python3 generate_docx_manual.py
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets the background color of a table cell (hex string e.g. 'EEF4F8')."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner margins (padding) for a cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="KARIS OS™ ABSOLUTE INVARIANT"):
    """Creates a beautifully styled callout box with a thick left navy border and light background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EEF4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=140)
    
    # Set borders: thick left border (#0B2545), none for top/bottom/right
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="0B2545"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"★ {title}\n")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph()  # spacing after table

def build_docx():
    doc = docx.Document()
    
    # Set standard page margins (1 inch around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set Normal style typography
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15
    
    # --- TITLE & SUBTITLE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("KARIS OS™ UNIFIED ENTERPRISE &\nDIGITAL ECONOMY PLATFORM")
    r_title.bold = True
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45) # Navy Blue
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Enterprise Architecture, Engineering Specification & Build Manual\nVersion 1.0.0-PROD-V1 | Flagship Production Vertical: KARIS FARM™")
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x13, 0x40, 0x74) # Slate Grey
    
    # Divider line
    p_div = doc.add_paragraph("─" * 55)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.runs[0].font.color.rgb = RGBColor(0x8C, 0x9B, 0xAD)
    
    # --- EXECUTIVE SUMMARY ---
    h1 = doc.add_heading("1. EXECUTIVE SUMMARY & PLATFORM VISION", level=1)
    h1.style.font.name = 'Arial'
    h1.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS OS™ is a modular, AI-native, event-driven Enterprise Operating System engineered to unify "
        "commerce, agriculture, retail, logistics, healthcare, mobility, lending, payments, and digital value exchange "
        "into one interoperable digital platform. Rather than building siloed, fragmented applications for disparate industries, "
        "KARIS OS™ provides an immutable, highly auditable operating foundation upon which unlimited vertical-specific solutions operate."
    )
    
    doc.add_paragraph(
        "This manual formalizes the complete 48-section architectural blueprint, translating conceptual vision directly into "
        "production-grade PostgreSQL DDL schemas, Draft-07 JSON Event specifications, double-entry Universal Ledger workflows, "
        "and concrete Python microservice orchestration. The flagship deployment, KARIS FARM™, demonstrates live digital "
        "traceability from smallholder farm registration across East Africa to retail checkout and automated M-Pesa settlement."
    )
    
    # --- SECTION 2: CORE RULES ---
    h1_2 = doc.add_heading("2. THE TEN ABSOLUTE PLATFORM INVARIANTS", level=1)
    h1_2.style.font.name = 'Arial'
    h1_2.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "To guarantee auditability, financial integrity, and strict economic boundaries across all multi-tenant organizations, "
        "every database trigger, microservice, and API route strictly enforces the following ten invariants without exception:"
    )
    
    add_callout_box(doc, "Rule 1: No Event → No State Change. All state mutations must originate from validated, timestamped events published to the Universal Event Bus.", "RULE 1: EVENT-DRIVEN INVARIANT")
    add_callout_box(doc, "Rule 2: No Payment → No Settlement. Merchant, supplier, and farmer settlements occur strictly after confirmed payment events are recorded.", "RULE 2: SETTLEMENT INVARIANT")
    add_callout_box(doc, "Rule 3: No Credit Approval → No Credit Purchase. Credit facilities mandate explicit AI-scored, lender-approved workflow events.", "RULE 3: CREDIT INVARIANT")
    add_callout_box(doc, "Rule 4: No Delivery Confirmation → No Rider Payment. Logistics payouts require cryptographically verified proof of delivery.", "RULE 4: LOGISTICS INVARIANT")
    add_callout_box(doc, "Rule 5: No Wallet Directly Edits Another Wallet. Every balance transfer must occur through: Ledger Engine → Wallet Engine → Settlement Engine.", "RULE 5: WALLET ISOLATION INVARIANT")
    add_callout_box(doc, "Rule 6: Everything Generates an Event. Every commercial transaction, identity verification, token grant, or inventory movement emits an immutable JSON event.", "RULE 6: AUDIT TRAIL INVARIANT")
    add_callout_box(doc, "Rule 7: Everything is Configurable. Business rules, reward ratios, workflow steps, and pricing parameters require zero source-code changes.", "RULE 7: DECLARATIVE CONFIGURATION")
    add_callout_box(doc, "Rule 8: Every Event is Timestamped. All records capture exact UTC timestamps with millisecond precision.", "RULE 8: TEMPORAL PRECISION")
    add_callout_box(doc, "Rule 9: Every Transaction is Immutable. Universal Ledger and Event Store tables enforce append-only database triggers. Corrections occur via reversing entries.", "RULE 9: IMMUTABILITY INVARIANT")
    add_callout_box(doc, "Rule 10: AI Assists; Humans Approve Configurable Decisions. AI generates predictive insights and risk scores, while high-impact decisions mandate human authorization.", "RULE 10: AI GOVERNANCE INVARIANT")
    
    # --- SECTION 3: ARCHITECTURE & DATA FLOW ---
    h1_3 = doc.add_heading("3. ENTERPRISE ARCHITECTURE & COMPONENT TOPOLOGY", level=1)
    h1_3.style.font.name = 'Arial'
    h1_3.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS OS™ separates transactional data, event logs, analytics, and AI knowledge into distinct, highly optimized tiers. "
        "The Universal Event Bus serves as the central nervous system connecting all microservices and business verticals."
    )
    
    # Create Table for Components
    data = [
        ("API Gateway & Security", "FastAPI, OAuth2, JWT, Rate Limiting", "Enforces Single Identity Principle, RBAC authorization, rate limiting, and request correlation IDs."),
        ("Universal Event Bus", "Event Store (PostgreSQL), Redis Pub/Sub", "Enforces Rule 1 & Rule 8. Captures every system action as an append-only JSON event with SHA-256 cryptographic hashes."),
        ("Universal Ledger Engine", "Double-Entry SQL, SHA-256 Chaining", "Enforces Rule 5 & Rule 9. Total Debits must equal Total Credits. Prevents database updates or deletes via PostgreSQL triggers."),
        ("Multi-Asset Wallet & Exchange Engine", "Isolated Wallets, Dynamic Rates", "Manages KES, KRT, Loyalty, Credit, and Investment balances. Executes mixed KES+KRT redemptions via Exchange Engine."),
        ("Enterprise AI Gateway", "Multi-Agent Orchestration (Rule 10)", "Coordinates Executive AI, Risk AI, Commerce AI, Logistics AI, and Agriculture AI for risk scoring and dynamic routing."),
        ("Vertical 1: KARIS FARM™", "Produce Batches, QR Traceability", "Digitizes farm registration, crop plans, harvest logs, and end-to-end QR code food safety traceability from farm to table."),
        ("Vertical 2: Retail POS & Supermarket", "Omnichannel POS, Barcode Scan", "Handles multi-branch POS cashier sessions, barcode scanning, dynamic pricing, and mixed KES+KRT M-Pesa checkouts."),
        ("Vertical 3: Eatery & Food Services", "Kitchen Display System (KDS)", "Manages digital menus, ingredient recipe deduction, order kitchen queues (PREPARING -> READY), and waste tracking."),
        ("Vertical 4: Delivery & Logistics", "AI Dispatcher, Escrow Payouts", "Enforces Rule 4. Matches riders via AI, tracks GPS proof of delivery, and releases escrow payouts upon completion."),
        ("Vertical 5: Healthcare & Medical", "EMR Notes, Prescriptions, CHVs", "Manages patient profiles, teleconsultations, EMR clinical notes, digital prescriptions, and CHV maternal health checks."),
        ("Vertical 6: Mobility & Ride-Hailing", "Drivers, Surge Pricing, Trips", "Matches drivers to passenger ride requests, calculates dynamic surge pricing via Mobility AI, and settles driver payouts."),
        ("Vertical 7: Investor & Lending", "Capital Pools, AI Credit Scoring", "Enforces Rule 3. Manages investment pools (Agriculture Growth Fund), ROI distribution, and AI-scored loan disbursements."),
        ("Vertical 8: Sales Force Automation", "Field Agents, Visits, Referrals", "Manages field representatives in Machakos County, logs customer visits, and auto-converts referrals (`100 KRT` bonus)."),
        ("Vertical 9: Loyalty & Incentives", "Promotional Campaigns, Grants", "Governs multi-stakeholder campaigns (`Harvest Double Points`), tracking earned points and KRT conversion."),
        ("Workflow State Machine Engine", "Declarative Approval Gates", "Enforces multi-step approval workflows (`CREDIT_APPROVAL_WORKFLOW`). AI auto-advances or escalates to human review (`Rule 10`)."),
        ("Dynamic Vertical Registry", "Industry Expansion Framework", "Enforces Section 35 (`Build once. Configure many. Scale infinitely`). Dynamically registers future industries (`Education`, `Tourism`)."),
        ("KARIS OS 2.0 Innovation Engine", "CBDC, Open Banking, EAC & ESG", "Enforces Section 48. Manages wholesale interbank CBDC-KES settlement, PSD2 consents, East African regional FX transfers (`KES->UGX`), and Scope 1/2/3 Carbon Footprint tracking awarding `KRT-GREEN` tokens."),
        ("WhatsApp Cloud Bot Engine", "Daraja & WhatsApp Business API", "Enables East African farmers & consumers to trace produce QR codes, check wallet balances, and ask grounded AI RAG farming questions right inside WhatsApp (`Rule 6`)."),
        ("CQRS Projections Engine", "Decoupled Event Read Models", "Enforces Section 37.3 (`APIs -> Events -> Read Models`). Asynchronously projects high-speed analytical summary dashboards without locking transactional write models."),
        ("Operational Governance & AML", "Multi-Tier KYC, SAR & KRA eTIMS", "Enforces Section 38.8 & 47. Executes multi-tier KYC (`TIER_3_ADVANCED`), flags AML structuring velocity checks (`SAR`), and issues digital KRA eTIMS tax invoices with exact VAT (`16%`)."),
        ("AI Predictive Forecasting", "Time-Series Stockout & Elasticity", "Enforces Section 27.4 & 13.4. Predicts SKU demand velocity, exact stockout dates across Machakos/Nairobi supermarket branches, and dynamic clearance pricing under `Rule 10`."),
        ("DLQ Self-Healing Engine", "Exponential Backoff Recovery", "Enforces Section 36.6. Captures transient event subscriber database timeouts into the Dead-Letter Queue (`DLQ`), auto-recovering events (`DLQ_EVENT_RECOVERED`) without data loss (`Rule 1 & 9`)."),
        ("Prometheus Observability Engine", "Plain-Text Telemetry /metrics", "Enforces Section 40.7 & 47.3. Exposes real-time Prometheus metrics (`karis_http_requests_total`, `karis_ledger_entries_total`, etc.) for instant Grafana scraping and SLA tracking."),
        ("Disaster Recovery & Event Replay", "PITR Snapshots & State Rebuild", "Enforces Section 37.5 & 44.4. Captures cryptographic point-in-time snapshots (`PITR`) to disk (`SHA-256 checksum`), and reconstructs exact wallet balances directly from replayable event sourcing logs."),
        ("API Gateway Security Middleware", "Rate Limiting & Correlation ID", "Enforces Section 36.1 & 46. Enforces Token Bucket rate limiting (`429 Too Many Requests`), correlation ID injection (`X-Correlation-ID`), CORS, and enterprise security headers."),
        ("AI Emergency Ambulance Dispatch", "Geodesic Proximity Matching", "Enforces Section 32.8 & 33.5. Matches nearest active ALS/BLS ambulances (`AMB-MACHAKOS-ALS-01`) using Haversine geodesic distance, checks 100% oxygen reserves, and dispatches in milliseconds."),
        ("AI POS Queue & Shrinkage Engine", "Congestion Alerts & Audits", "Enforces Section 20.3 & 30.5. Monitors cashier lines (`active_customers > 6`), auto-issuing `QUEUE_CONGESTION_DETECTED` to open Express Lanes (`POS-MLO-02`), and investigates physical recount discrepancies."),
        ("Future Industries Suite", "Education, Tourism & Real Estate", "Enforces Section 35.3. Inherits core shared services across `KARIS Edu-Pay` tuition fee installments (`150 KRT-EDU` bonus), `KARIS Safari` eco-lodge bookings (`KRT-GREEN`), and `KARIS Prop-Share` fractional property syndication."),
        ("AI CRM Churn & LTV Engine", "Win-Back Retention Campaigns", "Enforces Section 23.4. Predicts exact Customer Lifetime Value (`LTV`), estimates churn risk (`82.5%`), and auto-dispatches **500 KRT loyalty retention grants (`Rule 6 & 7`)** to win back at-risk users."),
        ("AI Fraud Detection Engine", "Geodesic Impossible Travel Check", "Enforces Section 38.6 & 27.1. Evaluates transaction velocity, impossible travel (`Machakos -> Mombasa in 3 mins`), and blacklisted devices, auto-freezing target wallets under `Rule 5 & 10`."),
        ("Call Center SLA Engine", "First Response & Escalation Sweep", "Enforces Section 24.4. Monitors SLA benchmarks (`2-hour First Response`), running automated sweeps to flag breached tickets and auto-assign them to senior supervisors (`TICKET_ESCALATED`)."),
        ("Omnichannel Portal Gateway", "Super App & Merchant Dashboard", "Enforces Section 31.2 & 31.3 (`One Account across 12 Verticals`). Aggregates single-account Super App profiles (`Amina Wanjiku`) and unified multi-store Merchant Portal summaries (`Machakos Hub`)."),
        ("ERP/SAP Accounting Sync Engine", "External Fiscal Batch Sync", "Enforces Section 36.5 & 43.2. Pushes verified double-entry ledger batches directly to `SAP S/4HANA Kenya` (`FY-2026-Q3`) and renders declarative notification templates (`SMS`, `WhatsApp`, `Email`)."),
        ("Parametric Crop Insurance & IoT", "Smart Irrigation & Claim Payout", "Enforces Section 34.4 & 28.5. Monitors IoT sensor telemetry (`soil_moisture < 20%`), auto-triggering smart irrigation (`2,500L water`) AND auto-settling **KES 50,000 claim payouts via Universal Ledger (`Rule 2 & 5`)**."),
        ("Operational Governance & API Keys", "SHA-256 Key Hashing & Tax Rules", "Enforces Section 43 & 47. Manages operational policies (`POL-GOV-RESERVE-RATIO-20PCT`), issues/rotates cryptographic API keys (`KARIS_LIVE_...`), and registers dynamic tax brackets (`16% VAT`, `5% WHT`)."),
        ("GDPR & Kenya DPA Privacy Engine", "Right-to-be-Forgotten Anonymizing", "Enforces Section 38.7 (`Consent & Data Export`). Compiles structured JSON personal data exports across 12 verticals, and anonymizes PII (`full_name -> ANONYMIZED_USER_...`) while leaving double-entry SHA-256 hash chains 100% intact (`Rule 9`)."),
        ("Chaos Engineering Resilience Suite", "Fault Injection & DLQ Drills", "Enforces Section 44.2 & 40.7. Simulates database slave connection timeouts and network latency during concurrent double-entry transfers, verifying our `DLQ Self-Healing Engine` auto-recovers (`DLQ_EVENT_RECOVERED`) with zero data loss."),
        ("Multi-Warehouse & Weather Dispatch", "Serial Crate Barcode & Storm AI", "Enforces Section 15 & 21.5. Tracks exact crate serial numbers (`SN-AVO-CRATE-2026-0001`) across warehouses (`Machakos/Mlolongo`), and auto-switches vehicle dispatches from motorcycles to covered refrigerated vans during heavy storms (`Surge 1.35x`)."),
        ("Automated API SDK Generator", "Python Async/Sync & TypeScript", "Enforces Section 46.2. Generates complete, type-hinted client packages (`karis_os_client.py`, `karis-os-sdk.ts`) wrapping all 12 verticals with correlation ID headers and JWT auth (`Rule 6`)."),
        ("Unified BI Executive Aggregator", "5-Domain C-Suite Intelligence", "Enforces Section 27.2 & 27.3. Compiles real-time executive briefings across `Executive Summary`, `Commerce Retail POS`, `Delivery Logistics`, `Finance Treasury Lending`, and `Healthcare EMR CHV`."),
        ("CI/CD Deployment Quality Gate", "Automated Release Verification", "Enforces Section 40.5 & 40.6. Evaluates build thresholds (`100% Pytest`, `Rule 9 clean`, `>1,500 ops/sec`). Issuing formal deployment authorization (`CICD_GATE_PASSED_AUTHORIZED`) or enforcing rollback."),
        ("Offline App & POS Sync Engine", "Optimistic Lock & Reconnect", "Enforces Section 41.5 & 20.2. Allows checkouts (`POS-MLO-01`) during network outages, executing optimistic lock verification and double-entry transfers upon reconnection (`OFFLINE_BATCH_SYNCHRONIZED`)."),
        ("Multi-Vendor Marketplace Engine", "Atomic Split-Commissions", "Enforces Section 14.3 & 17.2. Processes combined multi-vendor carts (`Farmer Kamau KES 5k + Coop KES 3k`), calculating split-commissions (`85% net, 15% fee`) across atomic multi-leg transfers (`Rule 2 & 5`)."),
        ("Regulatory Compliance Engine", "Multi-Jurisdictional Audits", "Enforces Section 35.4 & 38.8. Scans multi-tier KYC records, SARs, eTIMS invoices, and SHA-256 hash chains to compile verifiable government inspection packages (`CENTRAL_BANK_AML_FIU_SUMMARY`)."),
        ("HSM & Biometric NFC Engine", "AES-256-GCM Cryptogram Tokens", "Enforces Section 41.4, 20.1 & 38.4. Issues 60-second one-time NFC/QR cryptograms (`NFC-TOKEN-...`) after Face ID biometric verification, settling contactless checkouts via Universal Double-Entry Ledger (`Rule 5 & 9`)."),
        ("Customer Loyalty Tier & Network Engine", "VIP Auto-Upgrades & Clearing", "Enforces Section 23.2 & 18. Evaluates rolling spend (`KES 100k -> PLATINUM_VIP` with 15% delivery rebate), and clears cross-merchant KRT redemptions (`KARIS FARM -> Eatery`) via inter-org double-entry transfers (`Rule 5`)."),
        ("Active-Active HA & Failover Routing", "Geographic Cluster Nodes", "Enforces Section 40.8 & 45.3. Tracks live heartbeats across `Nairobi Main`, `Machakos Edge`, and `Mombasa DR` clusters. Executes instant geographic failovers (`FAILOVER_EXECUTED_TO_MACHAKOS_EDGE`) with 100% ledger continuity."),
        ("Mobile Passkeys & Push Notification", "FIDO2 WebAuthn & APNS/FCM", "Enforces Section 41.2, 41.3 & 26.5. Verifies hardware FIDO2 passkeys (`PASSKEY-CRED-01`), triggering gamified **250 KRT Security Champion bonuses (`Rule 6`)** every 10th login, and delivering live APNS/FCM mobile push alerts."),
        ("Double-Entry Ledger Recon Sweeps", "Mathematical Sum vs Wallets", "Enforces Section 37.4 & 10.4. Sweeps every multi-asset wallet against exact double-entry transaction debits/credits (`Sum(Credits)-Sum(Debits)`). Corrects drift exclusively via cryptographic Reversing Entries (`Rule 9`)."),
        ("Kubernetes Container Autoscaling", "Autonomous HPA Pod Scale-Out", "Enforces Section 40.3 & 40.4. Evaluates request velocity (`ops/sec`) and CPU load (`88.5%`). During traffic surges, dynamically scales microservice pod replicas from `4 -> 16 replicas` (`PODS_SCALED_OUT_TRAFFIC_SURGE`)."),
        ("Unified Escrow & Dispute Clearing", "Partial Refund Double-Entry Split", "Enforces Section 31.1 & 11.2. Holds multi-party checkouts (`KES 250k`) in secure escrow (`ESCROW-HOLD-01`), resolving produce quality disputes by executing atomic double-entry splits via Universal Ledger (`Rule 2, 5 & 9`)."),
        ("AI Supply Chain Bottleneck Engine", "Dynamic Highway Route Bypass", "Enforces Section 27.4 & 13.4. Monitors transit delays across East African distribution corridors (`Machakos-Nairobi Highway`). When backlogs surge (`>6h, >500 crates`), auto-calculates dynamic bypass routes under `Rule 10`."),
        ("Declarative Tax Holiday Override", "Statutory Tariff Exemption Engine", "Enforces Section 43.1 & 43.2. Allows statutory exemptions (`0% VAT during planting season for certified farmers`) to dynamically override standard KRA eTIMS invoice calculation without modifying source code (`Rule 7`)."),
        ("POWER BOT X Prediction Economy", "WhatsApp-First Agent Ecosystem", "Enforces Section 49 / Vertical 14 across all 5 layers (`Experience`, `AI`, `Economy`, `Growth`, `Infrastructure`). Features WhatsApp Status poster generation, **AI Prediction Copilot** (`improves decision quality without promising outcomes`), non-financial reputation graph (`unlocks experiences, not game outcomes`), double-entry prediction escrow (`Rule 9`), and immediate KRT winning redemptions across **KARIS Eatery** and **KARIS FARM**."),
        ("KARIS ENERGY & Smart Solar Grid", "PAYG Solar & P2P Energy Trading", "Enforces Section 50 / Vertical 15. Digitizes PAYG solar home systems and smart irrigation pumps (`SOLAR-PUMP-MACHAKOS-01`). Streams daily IoT smart meter telemetry (`Rule 6`), automatically minting **KRT-JOULE surplus rewards** when battery > 90% and power feeds back into the community grid (`Rule 7 & 9`). Enables P2P microgrid energy trading via exact double-entry accounting (`Rule 5`)."),
        ("PalPlus Hosted Payment Links Engine", "Active Temporary Checkout URL", "Enforces Section 51 / Section 34.5. Integrates universal hosted checkout URLs (`https://link.palpluss.com/6e8de0bc-1284-4bba-a5de-f886665bf18f`). Allows merchants, prediction leagues, and PAYG solar invoices to generate or attach hosted payment links. Reconciles M-Pesa Express webhooks (`Rule 2 & Rule 9`), triggering exact double-entry KES settlement and loyalty reward token minting (`Rule 5, 6 & 7`)."),
        ("KARIS Innovation Expansion Suite", "Pharma-Trace / Prop-Share / Edu-Pay", "Enforces Section 52 / Verticals 16, 17, 18. **Pharma-Trace** monitors cold-chain drug deliveries (`< 8°C`) and auto-locks breached batches (`Rule 1 & 6`). **Prop-Share** manages commercial real estate syndication (`Section 35.3`) and executes atomic double-entry monthly rental dividend distributions (`Rule 5 & Rule 9`). **Edu-Pay** digitizes student tuition fee installments, awarding **+150 KRT-EDU scholarship bonus tokens** (`Rule 7 & Rule 9`) upon payment verification."),
        ("White-Label Customization Engine", "Dynamic Multi-Tenant Branding", "Enforces Section 53 / Frontier D. Allows commercial partners (`Safaricom M-Pesa Enterprise`, `Equity Bank Fintech Hub`, `PalPlus Global Checkout OS`) to dynamically reconfigure platform metadata, color palettes (`#10B981 Green`), active payment links (`6e8de0bc...`), and default currencies (`KES/USD/UGX`) without modifying core operating system kernel code (`Rule 7 & Rule 9`)."),
        ("Karis Loop™ Social Intelligence Layer", "7 Graphs & Shoppable Checkouts", "Enforces Section 54 / Vertical 19. Unifies 7 interconnected graphs (`Social, Interest, Creator, Business, Community, Knowledge, Commerce`), 14 content types, multi-priority feed ranking (`Rule 7`), and creator KRT tipping (`Rule 5 & Rule 9 double entry`). Features **AI Content Assistant (`Swahili/Sheng RAG captions & moderation per Rule 10`)** and instant shoppable video checkouts (`One Wallet Economy`)."),
        ("KARIS Academy™ Educational Ecosystem", "Concept Graphs & Rule 10 Copilot", "Enforces Section 55 / Vertical 20. Organizes learning into interconnected concept graphs (`Linear Algebra -> Python -> AI RAG`), generates 14 original educational artifacts under **Rule 10 (`DRAFT_PENDING_EDUCATOR_APPROVAL`)**, records immutable academic transcripts (`Rule 9`), and automatically mints **`+250.00 KRT-EDU` utility reward tokens (`Rule 7 & Rule 9`)** directly to student wallets upon concept mastery.")
    ]

    tbl_comp = doc.add_table(rows=len(data) + 1, cols=3)
    tbl_comp.style = 'Table Grid'
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = tbl_comp.cell(row_idx, col_idx)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            p.add_run(text)
            
    doc.add_paragraph() # Spacing
    
    # --- SECTION 4: FLAGSHIP VERTICAL & SIMULATION ---
    h1_4 = doc.add_heading("4. FLAGSHIP VERTICAL: KARIS FARM™ & EXECUTION PROOF", level=1)
    h1_4.style.font.name = 'Arial'
    h1_4.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS FARM™ is the first production deployment of KARIS OS™, demonstrating full end-to-end agricultural value chain "
        "traceability. When Farmer John Kamau logs a 1,000 KG harvest of Grade-A Hass Avocados in Machakos County, the system "
        "automatically issues a Traceability QR Code (e.g., KARIS-TRACE-QR-0BA22758)."
    )
    doc.add_paragraph(
        "When Customer Amina Wanjiku purchases 50 KG of these Avocados for KES 7,500 via M-Pesa, the event bus triggers "
        "automated double-entry ledger settlement and awards a 5% KRT loyalty grant (375 KRT) directly from the Treasury Reward Pool. "
        "All data transformations occur in under 350 milliseconds and pass all cryptographic verification checks."
    )
    
    # Code block representing verification output
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.4)
    r_code = p_code.add_run(
        "[STEP 7] Verifying Final Wallet Balances after Ledger Settlement...\n"
        "  -> Customer KES Balance: KES 42,500.00 (Debited KES 7,500)\n"
        "  -> Farmer KES Balance:   KES 7,500.00 (Credited KES 7,500 via Ledger)\n"
        "  -> Customer KRT Balance: 375.00 KRT (Earned 5% Loyalty Reward)\n"
        "  -> Treasury KRT Pool:    499,625.00 KRT\n\n"
        "[STEP 8] Audit Proof: Universal Ledger Entries (SHA-256 Cryptographically Hashed)...\n"
        "  [1] AssetType.KES Transfer of 7,500.00 KES | Hash: 67e615c4c2cbd7ed51b1545cf...\n"
        "  [2] AssetType.KRT Transfer of 375.00 KRT   | Hash: 83902b1e356db10c36ff8a532..."
    )
    r_code.font.name = 'Courier New'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph() # Spacing

    # --- SECTION 5: POWER BOT X AUTONOMOUS AI PREDICTION ECONOMY ---
    h1_5 = doc.add_heading("5. SECTION 49 :: POWER BOT X AUTONOMOUS AI PREDICTION ECONOMY", level=1)
    h1_5.style.font.name = 'Arial'
    h1_5.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "Power BOT X operates as an integrated core vertical (`Section 49 / Vertical 14`) within KARIS OS™. "
        "It is strictly not a betting website, chatbot, or standalone mobile app—it is a self-improving digital prediction economy where AI, "
        "KRT tokens, autonomous agents, merchants, and users continuously create value across five interconnected layers:"
    )
    doc.add_paragraph(
        "1. Experience Layer (WhatsApp-First): AI automatically generates localized Swahili/Sheng/English Status posters, animated match cards, "
        "countdown timers, QR deep links (`wa.me`), and audio voice notes. Discovery and engagement return directly to WhatsApp.\n"
        "2. AI Layer (7 Intelligence Engines): The AI Prediction Copilot improves user decision quality by explaining tactical form, fixture congestion, "
        "and 95% confidence intervals without promising match outcomes (`Rule 10`). The Living AI Content Engine ensures no two agents receive identical marketing kits.\n"
        "3. Economy Layer (Dynamic KRT): Every deposit converts to KRT (`1 KES = 1 KRT`). KRT powers prediction entries, 10% agent commissions, and "
        "immediate redemption at KARIS Eatery and KARIS FARM (`One Wallet Economy`).\n"
        "4. Growth Layer (Social & Agents): Features private and county leagues (`Machakos vs Nairobi`) alongside a non-financial Reputation Graph "
        "that rewards fair participation (`Reputation unlocks experiences, not game outcomes`).\n"
        "5. Infrastructure Layer (Kernel Sync): Shares One Identity, Multi-Asset Wallets (`Rule 5`), and double-entry SHA-256 hash chaining (`Rule 9`), "
        "monitored by a self-evolving Digital Twin simulation requiring explicit C-suite RBAC approval (`Rule 10`)."
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 6: KARIS ENERGY & SMART SOLAR GRID ---
    h1_6 = doc.add_heading("6. SECTION 50 :: KARIS ENERGY & SMART SOLAR GRID™ (VERTICAL 15)", level=1)
    h1_6.style.font.name = 'Arial'
    h1_6.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS ENERGY & SMART SOLAR GRID™ operates as Vertical 15 (`Section 50`) of KARIS OS™, digitizing Pay-As-You-Go (PAYG) solar home "
        "systems, battery banks, and smart irrigation pumps across rural East Africa and urban microgrids. It integrates across 4 core capabilities:\n"
        "1. PAYG Solar Unit Registration & Token Unlocking: Farmers purchase solar irrigation pumps (`SOLAR-PUMP-MACHAKOS-01`) unlocked on daily "
        "schedules using KRT tokens (`Rule 2 & 9`). Unlocking fees transfer directly from the farmer's KRT wallet to the energy merchant wallet.\n"
        "2. IoT Smart Meter Telemetry (`Rule 6 & Rule 8`): Daily metrics log exact generation (`6.85 kWh`), consumption (`4.45 kWh`), and battery voltage.\n"
        "3. Automated Green Token Minting (`KRT-JOULE` / `KRT-GREEN`): When surplus solar generation (`microgrid_feed_in_kwh > 0`) feeds back into "
        "the community microgrid, the Universal Ledger Engine automatically debits the Treasury Reserve Pool and credits the farmer's wallet (`Rule 7 & 9`).\n"
        "4. Peer-to-Peer Microgrid Trading: Enables community buildings (`Health Clinics`) to purchase surplus solar electricity from neighboring "
        "farmers (`10 kWh at 12 KRT/kWh`) via atomic double-entry ledger settlement (`Rule 5 & Rule 9`)."
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 7: PALPLUS HOSTED PAYMENT LINKS ENGINE ---
    h1_7 = doc.add_heading("7. SECTION 51 :: PALPLUS & HOSTED PAYMENT LINKS ENGINE", level=1)
    h1_7.style.font.name = 'Arial'
    h1_7.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS OS™ integrates universal hosted checkout URLs and temporary payment links directly into our double-entry accounting kernel (`Section 51 / Rule 2 & 9`). "
        "Our active temporary payment link (`https://link.palpluss.com/6e8de0bc-1284-4bba-a5de-f886665bf18f` — Link ID: 6e8de0bc-1284-4bba-a5de-f886665bf18f) "
        "enables any consumer, smallholder farmer, or merchant to complete instant KES fiat checkouts via M-Pesa Express, card, or bank transfers right from WhatsApp or mobile browsers.\n\n"
        "When a customer completes payment on the hosted PalPlus checkout link, our secure webhook endpoint (`/api/v1/payment-links/webhook/palplus`) "
        "verifies the transaction (`PALPLUS-RC-99021`), publishes `PAYMENT_LINK_CHECKOUT_COMPLETED`, debits the customer KES account, credits the supplier KES account (`Rule 5 & 9`), "
        "and triggers our Rule Engine to automatically mint a 5% KRT loyalty reward (`Rule 7`). Zero manual reconciliation or human intervention is required."
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 8: KARIS INNOVATION EXPANSION SUITE ---
    h1_8 = doc.add_heading("8. SECTION 52 :: KARIS INNOVATION EXPANSION SUITE (VERTICALS 16, 17, 18)", level=1)
    h1_8.style.font.name = 'Arial'
    h1_8.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS INNOVATION EXPANSION SUITE (`Section 52`) demonstrates zero-code kernel expansion (`Rule 7 & Section 35`) across 3 high-impact industry domains:\n"
        "1. KARIS HEALTH PHARMA-TRACE™ (`Vertical 16`): Digitizes cold-chain pharmaceutical tracking (`BATCH-PHARMA-INS-2026`) and IoT temperature telemetry (`< 8°C`). "
        "When transit temperature exceeds 8.0°C, the system automatically fires `PHARMA_COLD_CHAIN_BREACH_DETECTED`, alerts Logistics AI, and locks the lot from retail dispensing (`Rule 1 & 6`).\n"
        "2. KARIS PROP-SHARE & FRAC-EQUITY™ (`Vertical 17`): Digitizes fractional commercial real estate syndication (`Section 35.3`). When monthly rental income (`KES 150,000`) "
        "arrives, UniversalLedgerEngine calculates and executes atomic double-entry dividend distributions (`Rule 5 & Rule 9 double-entry accounting`) straight into shareholder KRT/KES wallets.\n"
        "3. KARIS EDU-PAY & CAMPUS GRID™ (`Vertical 18`): Manages student tuition fee installment schedules (`Rule 3`) across technical colleges. "
        "Reconciles payments, debits student KES, credits college KES (`Rule 9`), and awards `+150 KRT-EDU` scholarship/cafeteria bonus tokens (`Rule 7`) upon verified checkout!"
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 9: WHITE-LABEL CUSTOMIZATION ENGINE ---
    h1_9 = doc.add_heading("9. SECTION 53 :: WHITE-LABEL CLIENT CUSTOMIZATION ENGINE (`Frontier D`)", level=1)
    h1_9.style.font.name = 'Arial'
    h1_9.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "Through our White-Label Customization Engine (`src/core/whitelabel_engine.py`), commercial partners license and brand the entire operating system "
        "across East Africa in seconds without modifying underlying double-entry accounting kernel code (`Rule 7 & Rule 9`).\n"
        "Pre-Seeded Commercial Licensing Profiles:\n"
        "1. SAFARICOM_MPESA_ENTERPRISE: `M-Pesa Enterprise & Digital Economy OS` (`#10B981` Green). Powers smallholder aggregation (`KARIS FARM`), POS checkouts, and student tuition plans (`Edu-Pay`).\n"
        "2. EQUITY_BANK_FINTECH_HUB: `Equity Digital Banking & Agri-Fintech OS` (`#8B0000` Red/Maroon). Agricultural input financing (`Rule 3`), PAYG solar pump lending (`KARIS ENERGY`), and regional CBDC clearing.\n"
        "3. PALPLUS_GLOBAL_CHECKOUT_OS: `PalPlus Universal Commerce & Checkout OS` (`#2563EB` Blue). Universal hosted checkout URLs (`link.palpluss.com/6e8de0bc...`) across 18 industry verticals."
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 10: KARIS LOOP SOCIAL INTELLIGENCE LAYER ---
    h1_10 = doc.add_heading("10. SECTION 54 :: KARIS LOOP™ SOCIAL INTELLIGENCE LAYER (VERTICAL 19)", level=1)
    h1_10.style.font.name = 'Arial'
    h1_10.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "Karis Loop™ operates as Vertical 19 (`Section 54`) of KARIS OS™, establishing a unified social intelligence, content creation, "
        "and digital payments economy across East Africa. It integrates across 4 core capabilities:\n"
        "1. 7 Interconnected Graphs: Balances `Social, Interest, Creator, Business, Community, Knowledge, and Commerce` graphs to deliver "
        "rich multi-priority feeds (`For You, Trending, Local Agri`) without relying solely on raw popularity metrics (`Rule 7`).\n"
        "2. Creator Economy & KRT Tipping (`Rule 5 & Rule 9 double entry`): Allows users to tip creators in KRT tokens (`50 KRT`) and purchase "
        "digital guides (`Tip Jar & Digital Product Store`) with exact cryptographic SHA-256 hash chaining intact.\n"
        "3. AI Content Assistant & Rule 10 Moderation: Auto-generates localized Swahili, Sheng, and English captions with optimal hashtag "
        "schedules (`Rule 6`). Continuously evaluates toxicity (`0.0 to 100.0`); if toxicity > 60, status locks at `FLAGGED_PENDING_HUMAN_REVIEW` "
        "requiring human community supervisor RBAC sign-off before post removal (`Rule 10`).\n"
        "4. Shoppable Video Checkouts (`One Wallet Economy`): Every short video or story post can link directly to a physical product (`Grade-A Avocados`), "
        "service (`Telemedicine`), or PAYG solar pump (`SOLAR-01`). Customers tap 'Buy Now' to check out instantaneously via KES/KRT double-entry settlement!"
    )

    doc.add_paragraph() # Spacing

    # --- SECTION 11: KARIS ACADEMY EDUCATIONAL ECOSYSTEM ---
    h1_11 = doc.add_heading("11. SECTION 55 :: KARIS ACADEMY™ EDUCATIONAL ECOSYSTEM (VERTICAL 20)", level=1)
    h1_11.style.font.name = 'Arial'
    h1_11.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    
    doc.add_paragraph(
        "KARIS Academy™ operates as Vertical 20 (`Section 55`) of KARIS OS™, establishing the world's most advanced AI-powered educational ecosystem. "
        "It integrates across 4 core capabilities:\n"
        "1. Knowledge Engine & Concept Graphs: Organizes learning into directional prerequisite concept pathways (`Linear Algebra -> Python -> AI RAG`) "
        "rather than isolated courses. Mappings support multiple frameworks (`KENYA_CBC, TVET, CAMBRIDGE_IGCSE, UNIVERSITY_DEGREE`).\n"
        "2. AI Education Engine (`Rule 10 Human Gate`): Auto-generates 14 original educational artifacts (`Lesson notes, Quizzes, Rubrics, Certificates`). "
        "Every generated item strictly locks at `DRAFT_PENDING_EDUCATOR_APPROVAL` (`Rule 10`) until reviewed, edited, and signed off by human faculty.\n"
        "3. Concept Mastery & KRT-EDU Utility Rewards: When a student passes final mastery evaluation (`score >= 85%`), the UniversalLedgerEngine "
        "records an immutable academic transcript (`Rule 9`) and automatically debits the Treasury Reward Pool to credit the student's KRT wallet with "
        "`+250.00 KRT-EDU` utility reward tokens (`Rule 5 & Rule 9 double-entry accounting`).\n"
        "4. Institutional Scholarship Disbursements: Institutional scholarship pools (`POOL-ACADEMY-SCHOLARSHIP-01`) disburse tuition and living stipends "
        "directly to student wallets via double-entry accounting (`Rule 5 & Rule 9`), ensuring 100% transparent educational financing across Africa."
    )

    h1_12 = doc.add_heading("12. SECTION 56 :: KARISFX™ GLOBAL FINANCIAL ECOSYSTEM & KRT ECONOMY (VERTICAL 21)", level=1)
    h1_12.style.font.name = 'Arial'
    h1_12.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    doc.add_paragraph(
        "KARISFX™ (`Section 56 / Vertical 21`) is the flagship global financial platform of KARIS OS™, powered by the KRT (Karis Token) Economy Layer. "
        "It unifies 16 global asset classes (`Forex, Stocks, Commodities, ETFs, Bonds, Futures, Options, Money Markets, Mutual Funds, Tokenized Assets, AI Portfolios, Social Trading, Copy Trading`) "
        "into one shared double-entry financial ecosystem (`Rule 5 & Rule 9`). Key pillars include:\n"
        "1. KRT Foundation & Multi-Asset Wallets: Every onboarded user automatically receives a KRT Wallet, multi-currency wallets (`KES, USD, EUR, GBP, Stablecoin`), "
        "a Rewards Wallet (`AssetType.KRT_REWARDS`), and a unique Treasury Reference (`TREASURY-ACCOUNT-REF`). All fee settlements and KRT movements record immutable double-entry hash chains (`Rule 9`).\n"
        "2. KRT Staking Modules (Up to 60% Fee Discounts): Staking KRT into lockup pools (`30, 90, 180, 365 days`) unlocks tiered APY (`12% - 25%`), VIP trading execution, AI priority, and **up to 60% trading fee discounts**.\n"
        "3. AI Economy Engine (`Rule 10`): Delivers 13 AI financial services (`Market Intelligence, Portfolio Optimization, Strategy Builder, Risk Analysis, Trade Journal, AI Mentor, etc.`). All outputs mandate Rule 10 human/trader review for actual capital deployment.\n"
        "4. Transparent Reward Engine & Wash-Trading Guardrails: Rewards 10 platform activities (`Trading, Learning, Strategies, Referrals, Bug Reports, etc.`). Enforces velocity checks and wash-trading detection (`Rule 1 & Rule 6`).\n"
        "5. Strategy Marketplace & Decentralized Governance: Allows creators to publish strategies/bots with 85/15 split double-entry settlement, while KRT tokenholders vote on platform proposals under Rule 10 AI impact analysis."
    )

    h1_13 = doc.add_heading("13. SECTION 57 :: COSMOX™ UNIVERSAL AI MARKETPLACE & KRT ECONOMY (VERTICAL 22)", level=1)
    h1_13.style.font.name = 'Arial'
    h1_13.style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    doc.add_paragraph(
        "COSMOX™ (`Section 57 / Vertical 22`) is an AI-powered universal marketplace where KRT is the native utility token and primary mechanism for rewards, loyalty, and internal value exchange. "
        "It operates seamlessly across local fiat currencies (`KES, USD, EUR`) and multi-channel payment gateways (`M-Pesa Daraja, PalPlus, Cards, Bank Transfers`). Key capabilities include:\n"
        "1. Universal Wallets: Onboarded users receive 6 unified wallets (`fiat, krt, rewards, escrow, merchant, driver`) supporting instant transfers, cashback, loyalty rewards, and merchant settlements (`Rule 5 & Rule 9`).\n"
        "2. AI Engine (7 Modules): Powers personalized shopping recommendations, elasticity-adjusted dynamic merchant pricing, inventory forecasting, multi-lingual translation (`SW, SHENG, EN, FR, AR`), fraud velocity detection, customer support, and governance advisory (`Rule 10`).\n"
        "3. Escrow Checkout & Strict Rule 4 Logistics Release: Buyer checkout locks KRT inside the Order Escrow Wallet (`Rule 2`). Driver delivery dispatch locks rider KRT bonus in Main Escrow (`Rule 4`). Upon `DELIVERY_CONFIRMED`, the engine strictly releases 88% to the seller, 12% to treasury minus **2% deflationary burn (`burn_krt`)**, +1.5% cashback to the buyer, and +25 KRT bonus to the driver (`Rule 4 & Rule 9`).\n"
        "4. Referral Network & Digital Services: Disburses transparent rewards across `INDIVIDUAL (+100 KRT)`, `MERCHANT (+500 KRT)`, and `DELIVERY_PARTNER (+250 KRT)` tiers (`Rule 1 & Rule 6`), while enabling developer AI tool/software checkouts with 85/15 split settlement.\n"
        "5. Multi-Signature Treasury & Tokenomics: High-value treasury requests (`> 100,000 KRT` or `$10,000 USD`) require explicit multi-signature RBAC sign-off (`Rule 10`) before double-entry disbursement."
    )

    # Save document
    doc.save("/home/user/karis_os_enterprise_architecture_and_build_manual_v1.docx")
    print("Document successfully created at /home/user/karis_os_enterprise_architecture_and_build_manual_v1.docx")

if __name__ == "__main__":
    build_docx()
